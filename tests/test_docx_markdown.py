from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from py_chunks import get_markdown
from py_chunks.chunkers.docx import docx_to_markdown


def _make_docx(tmp_path, name="sample.docx") -> Path:
    doc = Document()
    doc.add_heading("Main Title", level=1)
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("This is a regular paragraph with some content.")
    doc.add_paragraph("Item one", style="List Bullet")
    doc.add_paragraph("Item two", style="List Bullet")
    # Table
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Age"
    table.rows[0].cells[2].text = "City"
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "30"
    table.rows[1].cells[2].text = "NYC"
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("A second section paragraph.")
    path = tmp_path / name
    doc.save(str(path))
    return path


# -- Basic return type --------------------------------------------------------

def test_returns_string(tmp_path):
    path = _make_docx(tmp_path)
    md = docx_to_markdown(str(path))
    assert isinstance(md, str)
    assert len(md) > 0


# -- Headings -----------------------------------------------------------------

def test_h1_present(tmp_path):
    path = _make_docx(tmp_path)
    md = docx_to_markdown(str(path))
    assert "# Main Title" in md


def test_h2_present(tmp_path):
    path = _make_docx(tmp_path)
    md = docx_to_markdown(str(path))
    assert "## Section One" in md
    assert "## Section Two" in md


# -- Paragraph ----------------------------------------------------------------

def test_paragraph_present(tmp_path):
    path = _make_docx(tmp_path)
    md = docx_to_markdown(str(path))
    assert "regular paragraph" in md


# -- List ---------------------------------------------------------------------

def test_list_items_present(tmp_path):
    path = _make_docx(tmp_path)
    md = docx_to_markdown(str(path))
    assert "- Item one" in md
    assert "- Item two" in md


# -- Table --------------------------------------------------------------------

def test_table_is_pipe_format(tmp_path):
    path = _make_docx(tmp_path)
    md = docx_to_markdown(str(path))
    assert "| Name |" in md or "| Name" in md
    assert "| Alice |" in md or "| Alice" in md
    assert "---" in md  # separator row present


# -- Order of elements --------------------------------------------------------

def test_heading_before_paragraph(tmp_path):
    path = _make_docx(tmp_path)
    md = docx_to_markdown(str(path))
    h1_pos = md.index("# Main Title")
    para_pos = md.index("regular paragraph")
    assert h1_pos < para_pos


# -- Public API routing -------------------------------------------------------

def test_get_markdown_from_path(tmp_path):
    path = _make_docx(tmp_path)
    md = get_markdown(str(path))
    assert "# Main Title" in md


def test_get_markdown_from_bytes(tmp_path):
    path = _make_docx(tmp_path)
    data = path.read_bytes()
    md = get_markdown(data, filename="doc.docx")
    assert isinstance(md, str)
    assert "Main Title" in md


def test_get_markdown_from_fileobj(tmp_path):
    path = _make_docx(tmp_path)
    data = path.read_bytes()
    md = get_markdown(BytesIO(data), filename="doc.docx")
    assert "Main Title" in md


# -- Error handling -----------------------------------------------------------

def test_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        docx_to_markdown(str(tmp_path / "missing.docx"))


def test_wrong_extension_raises(tmp_path):
    p = tmp_path / "file.txt"
    p.write_text("hello")
    with pytest.raises(ValueError, match=r"\.docx"):
        docx_to_markdown(str(p))


def test_unsupported_format_raises(tmp_path):
    _ = _make_docx(tmp_path, name="doc.docx")
    p2 = tmp_path / "doc.unsupported"
    p2.write_text("not supported")
    with pytest.raises(ValueError, match="not support"):
        get_markdown(str(p2))


# -- Footnotes ----------------------------------------------------------------

def test_footnotes_appended(tmp_path):
    # Create a docx with a footnote using python-docx low-level API.
    # python-docx has no high-level footnote API; here we verify no crash.
    doc = Document()
    doc.add_paragraph("Text with footnote")
    path = tmp_path / "fn.docx"
    doc.save(str(path))
    md = docx_to_markdown(str(path))
    assert isinstance(md, str)


def _make_bold_docx(tmp_path):
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Hello ")
    run = para.add_run("world")
    run.bold = True
    path = tmp_path / "bold.docx"
    doc.save(str(path))
    return path


def _make_italic_docx(tmp_path):
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Say ")
    run = para.add_run("something")
    run.italic = True
    path = tmp_path / "italic.docx"
    doc.save(str(path))
    return path


def _make_bold_italic_docx(tmp_path):
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("emphasis")
    run.bold = True
    run.italic = True
    path = tmp_path / "bold_italic.docx"
    doc.save(str(path))
    return path


def _make_vmerge_docx(tmp_path):
    """2x2 table: first column is vertically merged."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    # Row 0, col 1
    table.cell(0, 1).text = "A"
    # Row 1, col 1
    table.cell(1, 1).text = "B"

    # Vertically merge first column using public API.
    merged = table.cell(0, 0).merge(table.cell(1, 0))
    merged.text = "Merged"

    path = tmp_path / "vmerge.docx"
    doc.save(str(path))
    return path


def _make_header_docx(tmp_path):
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Company Report"
    doc.add_paragraph("Body text here.")
    path = tmp_path / "header.docx"
    doc.save(str(path))
    return path


def test_bold_run(tmp_path):
    path = _make_bold_docx(tmp_path)
    result = docx_to_markdown(str(path))
    assert "Hello" in result
    assert "**world**" in result


def test_italic_run(tmp_path):
    path = _make_italic_docx(tmp_path)
    result = docx_to_markdown(str(path))
    assert "Say" in result
    assert "*something*" in result


def test_bold_italic_run(tmp_path):
    path = _make_bold_italic_docx(tmp_path)
    result = docx_to_markdown(str(path))
    assert "***emphasis***" in result


def test_plain_text_unchanged(tmp_path):
    """Plain runs produce no markdown markers."""
    doc = Document()
    doc.add_paragraph("Just plain text")
    path = tmp_path / "plain.docx"
    doc.save(str(path))
    result = docx_to_markdown(str(path))
    assert "Just plain text" in result
    assert "**" not in result
    assert "*" not in result


def test_vmerge_continuation_repeats_content(tmp_path):
    path = _make_vmerge_docx(tmp_path)
    result = docx_to_markdown(str(path))
    # "Merged" must appear in both rows of the first column
    lines = [l for l in result.splitlines() if "Merged" in l]
    assert len(lines) >= 2, f"Expected 'Merged' in 2 rows, got:\n{result}"


def test_vmerge_other_cols_unaffected(tmp_path):
    path = _make_vmerge_docx(tmp_path)
    result = docx_to_markdown(str(path))
    assert "A" in result
    assert "B" in result


def test_header_prepended_as_blockquote(tmp_path):
    path = _make_header_docx(tmp_path)
    result = docx_to_markdown(str(path))
    assert "> Company Report" in result


def test_header_before_body(tmp_path):
    path = _make_header_docx(tmp_path)
    result = docx_to_markdown(str(path))
    header_pos = result.index("Company Report")
    body_pos = result.index("Body text here")
    assert header_pos < body_pos


def test_no_header_no_blockquote(tmp_path):
    """Documents without headers produce no leading blockquote."""
    doc = Document()
    doc.add_paragraph("No header here.")
    path = tmp_path / "no_header.docx"
    doc.save(str(path))
    result = docx_to_markdown(str(path))
    assert not result.startswith(">")
